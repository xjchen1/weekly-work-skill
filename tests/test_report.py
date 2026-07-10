import json
import os
import subprocess
import sys
import tempfile
import unittest
from datetime import date
from pathlib import Path
from unittest import mock

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt


SKILL_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(SKILL_ROOT))

from scripts import report, weekly_store


TEMPLATE_PATH = SKILL_ROOT / "assets" / "weekly-report-template.docx"
REPORT_SCRIPT = SKILL_ROOT / "scripts" / "report.py"


def add_entry(root, work_date, project, **sections):
    payload = {
        "date": work_date,
        "project": project,
        "ai_source": "codex",
        "daily_work": [],
        "literature": [],
        "research_progress": [],
        "research_outputs": [],
        "unresolved": [],
    }
    payload.update(sections)
    return weekly_store.add_entry(root, payload)


def load_doc(path):
    return Document(str(path))


def paragraph_texts(doc):
    return [paragraph.text for paragraph in doc.paragraphs]


def run_cli(*args, cwd=SKILL_ROOT):
    return subprocess.run(
        [sys.executable, str(REPORT_SCRIPT), *args],
        cwd=str(cwd),
        capture_output=True,
        text=True,
    )


class TemplateTests(unittest.TestCase):
    def setUp(self):
        self.tempdir = tempfile.TemporaryDirectory()
        self.root = Path(self.tempdir.name)

    def tearDown(self):
        self.tempdir.cleanup()

    def test_create_template_writes_exact_page_geometry_and_style_tokens(self):
        target = self.root / "weekly-report-template.docx"

        created = report.create_template(target)

        self.assertEqual(created, target)
        doc = load_doc(target)
        section = doc.sections[0]

        self.assertAlmostEqual(section.page_width.inches, 8.27, places=2)
        self.assertAlmostEqual(section.page_height.inches, 11.69, places=2)
        self.assertEqual(section.top_margin, Cm(2.54))
        self.assertEqual(section.bottom_margin, Cm(2.54))
        self.assertEqual(section.left_margin, Cm(2.54))
        self.assertEqual(section.right_margin, Cm(2.54))
        self.assertAlmostEqual(section.header_distance.inches, 0.492, places=3)
        self.assertAlmostEqual(section.footer_distance.inches, 0.492, places=3)

        normal = doc.styles["Normal"]
        title = doc.styles["Title"]
        heading1 = doc.styles["Heading 1"]
        heading2 = doc.styles["Heading 2"]
        work_item = doc.styles["Work Item"]

        self.assertEqual(normal.font.size, Pt(11))
        self.assertEqual(normal.font.name, "Arial")
        self.assertEqual(normal.paragraph_format.space_after, Pt(6))
        self.assertEqual(normal.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertAlmostEqual(normal.paragraph_format.line_spacing, 1.5)

        self.assertEqual(title.font.size, Pt(16))
        self.assertTrue(title.font.bold)
        self.assertEqual(title.font.color.rgb.__str__(), "000000")
        self.assertEqual(title.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(title.paragraph_format.space_after, Pt(10))

        self.assertEqual(heading1.font.size, Pt(14))
        self.assertTrue(heading1.font.bold)
        self.assertEqual(heading1.font.color.rgb.__str__(), "1F4D78")
        self.assertEqual(heading1.paragraph_format.space_before, Pt(14))
        self.assertEqual(heading1.paragraph_format.space_after, Pt(7))
        self.assertTrue(heading1.paragraph_format.keep_with_next)

        self.assertEqual(heading2.font.size, Pt(12))
        self.assertTrue(heading2.font.bold)
        self.assertEqual(heading2.font.color.rgb.__str__(), "1F4D78")
        self.assertEqual(heading2.paragraph_format.space_before, Pt(8))
        self.assertEqual(heading2.paragraph_format.space_after, Pt(4))
        self.assertTrue(heading2.paragraph_format.keep_with_next)

        self.assertEqual(work_item.paragraph_format.left_indent, Inches(0.25))
        self.assertEqual(work_item.paragraph_format.space_after, Pt(0))
        self.assertAlmostEqual(work_item.paragraph_format.line_spacing, 1.5)

        self.assertEqual(paragraph_texts(doc), [])
        self.assertEqual(len(doc.tables), 0)


class GenerateReportTests(unittest.TestCase):
    def setUp(self):
        self.tempdir = tempfile.TemporaryDirectory()
        self.root = Path(self.tempdir.name)
        self.template = self.root / "template.docx"
        report.create_template(self.template)

    def tearDown(self):
        self.tempdir.cleanup()

    def test_empty_week_raises_lookup_error_and_does_not_create_output(self):
        output = self.root / "out.docx"

        with self.assertRaises(LookupError):
            report.generate_report(
                self.root, date(2026, 6, 30), self.template, output_path=output
            )

        self.assertFalse(output.exists())

    def test_generate_report_uses_default_filename_and_expected_title(self):
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            daily_work=["Implemented template"],
        )
        add_entry(
            self.root,
            "2026-07-02",
            "beta",
            daily_work=["Verified report"],
        )

        output = report.generate_report(self.root, date(2026, 6, 30), self.template)

        self.assertEqual(
            output,
            self.root / "2026-06-29至2026-07-05周报.docx",
        )
        doc = load_doc(output)
        self.assertEqual(doc.paragraphs[0].text, "周报（06/29-07/05）")

    def test_fixed_sections_appear_once_in_order_and_daily_dates_are_sorted(self):
        add_entry(
            self.root,
            "2026-07-02",
            "beta",
            daily_work=["Second day"],
        )
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            daily_work=["First day"],
        )

        output = report.generate_report(self.root, date(2026, 6, 30), self.template)
        texts = [text for text in paragraph_texts(load_doc(output)) if text]

        self.assertEqual(texts[0], "周报（06/29-07/05）")
        self.assertLess(texts.index("1. 每日工作清单"), texts.index("2. 阅读文献情况"))
        self.assertLess(texts.index("2. 阅读文献情况"), texts.index("3. 研究进展情况"))
        self.assertLess(texts.index("3. 研究进展情况"), texts.index("4. 科研成果情况"))
        self.assertEqual(texts.count("1. 每日工作清单"), 1)
        self.assertEqual(texts.count("2. 阅读文献情况"), 1)
        self.assertEqual(texts.count("3. 研究进展情况"), 1)
        self.assertEqual(texts.count("4. 科研成果情况"), 1)
        self.assertLess(texts.index("06/30 星期二"), texts.index("07/02 星期四"))

    def test_empty_categories_render_as_wu(self):
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            daily_work=["Kept the week alive"],
        )

        output = report.generate_report(self.root, date(2026, 6, 30), self.template)
        texts = paragraph_texts(load_doc(output))
        self.assertIn("无", texts)
        self.assertEqual(texts.count("无"), 3)

    def test_literature_progress_and_outputs_render_content_and_unresolved_prefix(self):
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            daily_work=["Implemented reporting"],
            literature=[
                {"title": "Paper A", "author": "Li", "note": "read"},
            ],
            research_progress=[
                {"topic": "Experiment", "status": "running", "summary": "Stable"},
            ],
            research_outputs=[
                {"title": "Poster", "result": "submitted"},
            ],
            unresolved=[
                {"topic": "Experiment", "question": "Need clarification"},
            ],
        )

        output = report.generate_report(self.root, date(2026, 6, 30), self.template)
        texts = paragraph_texts(load_doc(output))
        self.assertIn("题目：Paper A；作者：Li；备注：read", texts)
        self.assertIn("主题：Experiment；状态：running；摘要：Stable", texts)
        self.assertIn("待确认：主题：Experiment；问题：Need clarification", texts)
        self.assertIn("题目：Poster；结果：submitted", texts)

    def test_canonical_duplicates_are_rendered_once_with_first_input_order(self):
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            literature=[
                {"title": "Paper A", "author": "Li"},
                {"title": "Paper A", "author": "Li"},
                {"title": "Paper B", "author": "Wang"},
            ],
        )

        output = report.generate_report(self.root, date(2026, 6, 30), self.template)
        texts = paragraph_texts(load_doc(output))
        self.assertEqual(texts.count("题目：Paper A；作者：Li"), 1)
        self.assertLess(
            texts.index("题目：Paper A；作者：Li"),
            texts.index("题目：Paper B；作者：Wang"),
        )

    def test_close_duplicates_keep_the_more_complete_item(self):
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            literature=[
                {"title": "Paper A"},
                {"title": "Paper A。", "author": "Li"},
                "完成 文献筛选。",
                "完成   文献筛选",
            ],
        )

        output = report.generate_report(self.root, date(2026, 6, 30), self.template)
        texts = paragraph_texts(load_doc(output))

        self.assertNotIn("题目：Paper A", texts)
        self.assertEqual(texts.count("题目：Paper A。；作者：Li"), 1)
        self.assertEqual(texts.count("完成 文献筛选。"), 1)
        self.assertNotIn("完成   文献筛选", texts)

    def test_similar_but_distinct_strings_are_not_merged(self):
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            literature=["Paper A", "Paper AB"],
        )

        output = report.generate_report(self.root, date(2026, 6, 30), self.template)
        texts = paragraph_texts(load_doc(output))

        self.assertIn("Paper A", texts)
        self.assertIn("Paper AB", texts)

    def test_identifier_values_remain_case_sensitive_during_report_deduplication(self):
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            literature=[
                {"title": "Artifact", "URL": "https://example.test/A"},
                {"title": "Artifact", "URL": "https://example.test/a"},
            ],
        )

        output = report.generate_report(self.root, date(2026, 6, 30), self.template)
        texts = paragraph_texts(load_doc(output))

        self.assertIn("题目：Artifact；URL：https://example.test/A", texts)
        self.assertIn("题目：Artifact；URL：https://example.test/a", texts)

    def test_unresolved_source_wins_cross_source_duplicate_without_reordering(self):
        duplicate_progress = {"topic": "Calibration", "status": "blocked"}
        duplicate_unresolved = {"status": "blocked", "topic": "Calibration"}
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            research_progress=[
                {"topic": "Progress before"},
                duplicate_progress,
                {"topic": "Progress after"},
            ],
            unresolved=[
                {"topic": "Unresolved before"},
                duplicate_unresolved,
                {"topic": "Unresolved after"},
            ],
        )

        output = report.generate_report(self.root, date(2026, 6, 30), self.template)
        texts = paragraph_texts(load_doc(output))
        duplicate_text = "待确认：主题：Calibration；状态：blocked"

        self.assertEqual(texts.count(duplicate_text), 1)
        self.assertNotIn("主题：Calibration；状态：blocked", texts)
        expected_order = [
            "主题：Progress before",
            "主题：Progress after",
            "待确认：主题：Unresolved before",
            duplicate_text,
            "待确认：主题：Unresolved after",
        ]
        self.assertEqual(
            [text for text in texts if text in expected_order],
            expected_order,
        )

    def test_partial_unresolved_item_does_not_hide_complete_progress(self):
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            research_progress=[
                {"topic": "模型A", "status": "blocked", "result": "基线完成"},
            ],
            unresolved=[{"status": "blocked"}],
        )

        output = report.generate_report(self.root, date(2026, 6, 30), self.template)
        texts = paragraph_texts(load_doc(output))

        self.assertIn("主题：模型A；结果：基线完成；状态：blocked", texts)
        self.assertIn("待确认：状态：blocked", texts)

    def test_atomic_validation_failure_preserves_existing_output(self):
        output = self.root / "existing.docx"
        output.write_text("keep-me", encoding="utf-8")
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            daily_work=["Prepared template"],
        )

        real_document = report.Document
        calls = []

        def fake_document(path):
            calls.append(str(path))
            if len(calls) == 2:
                raise ValueError("broken docx")
            return real_document(path)

        with mock.patch("scripts.report.Document", side_effect=fake_document):
            with self.assertRaises(ValueError):
                report.generate_report(
                    self.root,
                    date(2026, 6, 30),
                    self.template,
                    output_path=output,
                )

        self.assertEqual(output.read_text(encoding="utf-8"), "keep-me")

    def test_generate_report_prints_absolute_path_via_cli(self):
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            daily_work=["Implemented CLI"],
        )
        template = self.root / "cli-template.docx"
        report.create_template(template)

        create = run_cli("create-template", "--output", str(self.root / "cli-created.docx"))
        self.assertEqual(create.returncode, 0)
        self.assertTrue(Path(create.stdout.strip()).is_absolute())
        self.assertTrue(Path(create.stdout.strip()).exists())

        generate = run_cli(
            "generate",
            "--root",
            str(self.root),
            "--week-of",
            "2026-06-30",
            "--template",
            str(template),
        )
        self.assertEqual(generate.returncode, 0)
        generated = Path(generate.stdout.strip())
        self.assertTrue(generated.is_absolute())
        self.assertTrue(generated.exists())

    def test_generate_cli_rejects_invalid_week_of_with_useful_error(self):
        result = run_cli(
            "generate",
            "--root",
            str(self.root),
            "--week-of",
            "not-a-date",
            "--template",
            str(self.template),
        )

        self.assertEqual(result.returncode, 2)
        self.assertIn("invalid date: not-a-date", result.stderr)

    def test_generate_cli_rejects_missing_template_with_useful_error(self):
        add_entry(
            self.root,
            "2026-06-30",
            "alpha",
            daily_work=["Prepared report"],
        )
        missing_template = self.root / "missing-template.docx"

        result = run_cli(
            "generate",
            "--root",
            str(self.root),
            "--week-of",
            "2026-06-30",
            "--template",
            str(missing_template),
        )

        self.assertEqual(result.returncode, 2)
        self.assertIn(str(missing_template), result.stderr)


class TemplateCliSmokeTests(unittest.TestCase):
    def test_module_imports_for_tdd_red_phase(self):
        self.assertTrue(hasattr(report, "create_template"))


if __name__ == "__main__":
    unittest.main()
