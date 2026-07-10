#!/usr/bin/env python3
import argparse
import json
import os
import re
import sys
import tempfile
import unicodedata
from datetime import date
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


if __package__:
    from . import weekly_store
else:
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    import weekly_store


TITLE_STYLE = "Title"
HEADING1_STYLE = "Heading 1"
HEADING2_STYLE = "Heading 2"
WORK_ITEM_STYLE = "Work Item"

BODY_ASCII_FONT = "Arial"
BODY_EAST_ASIA_FONT = "Songti SC"
HEADING_EAST_ASIA_FONT = "Heiti SC"
TITLE_EAST_ASIA_FONT = "Songti SC"
HEADING_COLOR = "1F4D78"
BLACK = "000000"
RELATED_IDENTITY_FIELDS = ("topic", "title", "name", "content", "task")

WEEKDAY_LABELS = "一二三四五六日"
FIELD_LABELS = {
    "title": "题目",
    "topic": "主题",
    "project": "项目",
    "author": "作者",
    "status": "状态",
    "summary": "摘要",
    "result": "结果",
    "note": "备注",
    "question": "问题",
    "source": "来源",
    "progress": "进展",
    "plan": "计划",
    "conclusion": "结论",
    "method": "方法",
    "date": "日期",
    "ai_source": "来源",
}


def _set_run_fonts(style, ascii_font: str, east_asia_font: str) -> None:
    style.font.name = ascii_font
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def _set_paragraph_style(
    style,
    *,
    ascii_font: str,
    east_asia_font: str,
    size_pt: int,
    bold: bool = False,
    color: str = BLACK,
    alignment: Optional[int] = None,
    before_pt: Optional[int] = None,
    after_pt: Optional[int] = None,
    line_spacing: Optional[float] = None,
    left_indent: Optional[float] = None,
    keep_with_next: Optional[bool] = None,
) -> None:
    _set_run_fonts(style, ascii_font, east_asia_font)
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)

    pf = style.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if before_pt is not None:
        pf.space_before = Pt(before_pt)
    if after_pt is not None:
        pf.space_after = Pt(after_pt)
    if line_spacing is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)
    if keep_with_next is not None:
        pf.keep_with_next = keep_with_next


def _remove_paragraph_border(style) -> None:
    p_pr = style.element.pPr
    if p_pr is None:
        return
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def _configure_page(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def _clear_body_content(doc) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _ensure_work_item_style(doc):
    if WORK_ITEM_STYLE in doc.styles:
        style = doc.styles[WORK_ITEM_STYLE]
    else:
        style = doc.styles.add_style(WORK_ITEM_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    _set_paragraph_style(
        style,
        ascii_font=BODY_ASCII_FONT,
        east_asia_font=BODY_EAST_ASIA_FONT,
        size_pt=11,
        bold=False,
        color=BLACK,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after_pt=0,
        line_spacing=1.5,
        left_indent=0.25,
    )
    return style


def _configure_styles(doc) -> None:
    normal = doc.styles["Normal"]
    _set_paragraph_style(
        normal,
        ascii_font=BODY_ASCII_FONT,
        east_asia_font=BODY_EAST_ASIA_FONT,
        size_pt=11,
        bold=False,
        color=BLACK,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after_pt=6,
        line_spacing=1.5,
    )

    title = doc.styles[TITLE_STYLE]
    title.base_style = doc.styles["Normal"]
    _set_paragraph_style(
        title,
        ascii_font=BODY_ASCII_FONT,
        east_asia_font=TITLE_EAST_ASIA_FONT,
        size_pt=16,
        bold=True,
        color=BLACK,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after_pt=10,
    )
    _remove_paragraph_border(title)

    heading1 = doc.styles[HEADING1_STYLE]
    heading1.base_style = doc.styles["Normal"]
    _set_paragraph_style(
        heading1,
        ascii_font=BODY_ASCII_FONT,
        east_asia_font=HEADING_EAST_ASIA_FONT,
        size_pt=14,
        bold=True,
        color=HEADING_COLOR,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        before_pt=14,
        after_pt=7,
        keep_with_next=True,
    )

    heading2 = doc.styles[HEADING2_STYLE]
    heading2.base_style = doc.styles["Normal"]
    _set_paragraph_style(
        heading2,
        ascii_font=BODY_ASCII_FONT,
        east_asia_font=HEADING_EAST_ASIA_FONT,
        size_pt=12,
        bold=True,
        color=HEADING_COLOR,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        before_pt=8,
        after_pt=4,
        keep_with_next=True,
    )

    _ensure_work_item_style(doc)


def _set_section_properties(doc) -> None:
    _configure_page(doc.sections[0])


def _build_template_document():
    doc = Document()
    _set_section_properties(doc)
    _configure_styles(doc)
    _clear_body_content(doc)
    return doc


def _atomic_save_document(doc, output_path: Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    temp_handle = None
    temp_path = None
    try:
        temp_handle = tempfile.NamedTemporaryFile(
            mode="wb",
            dir=str(output_path.parent),
            prefix="." + output_path.stem + ".",
            suffix=".docx",
            delete=False,
        )
        temp_path = Path(temp_handle.name)
        temp_handle.close()
        doc.save(str(temp_path))
        Document(str(temp_path))
        os.replace(str(temp_path), str(output_path))
        temp_path = None
        return output_path
    finally:
        if temp_handle is not None:
            try:
                temp_handle.close()
            except Exception:
                pass
        if temp_path is not None:
            try:
                temp_path.unlink()
            except FileNotFoundError:
                pass


def create_template(path: Path) -> Path:
    path = Path(path)
    doc = _build_template_document()
    return _atomic_save_document(doc, path)


def _weekday_label(day: date) -> str:
    return WEEKDAY_LABELS[day.weekday()]


def _week_range_label(week_of: date) -> str:
    start, end = weekly_store.week_bounds(week_of)
    return start, end, "周报（%s-%s）" % (start.strftime("%m/%d"), end.strftime("%m/%d"))


def _default_output_path(root: Path, week_of: date) -> Path:
    start, end = weekly_store.week_bounds(week_of)
    return Path(root) / ("%s至%s周报.docx" % (start.isoformat(), end.isoformat()))


def _normalized_text(value: str, field_name: Optional[str] = None) -> str:
    normalized = re.sub(r"\s+", " ", unicodedata.normalize("NFC", value).strip())
    if field_name is not None and weekly_store._is_identifier_like_field_name(field_name):
        return normalized
    stripped = normalized.rstrip(weekly_store.TERMINAL_SENTENCE_PUNCTUATION)
    return (stripped or normalized).casefold()


def _canonical_key(value: Any, field_name: Optional[str] = None):
    if isinstance(value, dict):
        return (
            "dict",
            tuple(
                sorted(
                    (key, _canonical_key(item, field_name=key))
                    for key, item in value.items()
                )
            ),
        )
    if isinstance(value, list):
        return ("list", tuple(_canonical_key(item, field_name=field_name) for item in value))
    if isinstance(value, tuple):
        return ("tuple", tuple(_canonical_key(item, field_name=field_name) for item in value))
    if isinstance(value, str):
        return ("scalar", _normalized_text(value, field_name=field_name))
    return ("scalar", value)


def _subsumes(container: Any, candidate: Any, field_name: Optional[str] = None) -> bool:
    if isinstance(container, str) and isinstance(candidate, str):
        container_text = _normalized_text(container, field_name=field_name)
        candidate_text = _normalized_text(candidate, field_name=field_name)
        return container_text == candidate_text
    if isinstance(container, dict) and isinstance(candidate, dict):
        if _canonical_key(container) == _canonical_key(candidate):
            return True
        has_matching_identity = any(
            key in container
            and key in candidate
            and _canonical_key(container[key], field_name=key)
            == _canonical_key(candidate[key], field_name=key)
            for key in RELATED_IDENTITY_FIELDS
        )
        if not has_matching_identity:
            return False
        return all(
            key in container
            and _subsumes(container[key], value, field_name=key)
            for key, value in candidate.items()
        )
    return _canonical_key(container, field_name) == _canonical_key(candidate, field_name)


def _deduplicate(items: Iterable[Any]) -> List[Any]:
    result = []
    for item in items:
        for index, existing in enumerate(result):
            if _subsumes(existing, item):
                break
            if _subsumes(item, existing):
                result[index] = item
                break
        else:
            result.append(item)
    return result


def _is_blank(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        return not value.strip()
    return False


def _render_value(value: Any) -> str:
    if isinstance(value, str):
        return value
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return str(value)
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, list):
        if not value:
            return "[]"
        if all(not isinstance(item, (dict, list, tuple)) for item in value):
            return "、".join(_render_value(item) for item in value)
    return json.dumps(value, ensure_ascii=False, separators=(",", ":"))


def _field_label(key: str) -> str:
    return FIELD_LABELS.get(key, key)


def _render_item_text(item: Any, unresolved: bool = False) -> str:
    if isinstance(item, dict):
        parts = []
        ordered_keys = []
        for preferred in ("topic", "title"):
            if preferred in item:
                ordered_keys.append(preferred)
        for key in sorted(item.keys()):
            if key not in ordered_keys:
                ordered_keys.append(key)
        for key in ordered_keys:
            value = item[key]
            if _is_blank(value):
                continue
            parts.append("%s：%s" % (_field_label(key), _render_value(value)))
        text = "；".join(parts)
    else:
        text = str(item)

    if unresolved:
        return "待确认：" + text
    return text


def _add_paragraph(doc, text: str, style: str):
    return doc.add_paragraph(text=text, style=style)


def _add_item_paragraph(doc, item: Any, style: str = WORK_ITEM_STYLE, unresolved: bool = False):
    paragraph = doc.add_paragraph(style=style)
    if isinstance(item, dict):
        text = _render_item_text(item, unresolved=unresolved)
        paragraph.add_run(text)
        return paragraph
    if unresolved:
        prefix = paragraph.add_run("待确认：")
        prefix.bold = True
        paragraph.add_run(str(item))
        return paragraph
    paragraph.add_run(str(item))
    return paragraph


def _group_entries_by_date(entries: Sequence[dict]) -> Dict[date, List[dict]]:
    grouped: Dict[date, List[dict]] = {}
    for entry in entries:
        work_day = weekly_store.normalize_date(entry["date"])
        grouped.setdefault(work_day, []).append(entry)
    return grouped


def _section_has_any(entries: Sequence[dict], field_names: Sequence[str]) -> bool:
    for entry in entries:
        for field_name in field_names:
            if entry.get(field_name):
                return True
    return False


def _render_daily_section(doc, entries: Sequence[dict]) -> None:
    heading = _add_paragraph(doc, "1. 每日工作清单", HEADING1_STYLE)
    grouped = _group_entries_by_date(entries)
    if not any(entry.get("daily_work") for entry in entries):
        _add_item_paragraph(doc, "无")
        return

    for work_day in sorted(grouped):
        _add_paragraph(
            doc,
            "%s 星期%s" % (work_day.strftime("%m/%d"), _weekday_label(work_day)),
            HEADING2_STYLE,
        )
        for entry in grouped[work_day]:
            project = str(entry.get("project", "")).strip()
            if project:
                _add_item_paragraph(doc, "项目：%s" % project)
            for item in entry.get("daily_work", []) or []:
                _add_item_paragraph(doc, item)


def _render_flat_section(
    doc,
    title: str,
    entries: Sequence[dict],
    field_name: str,
    *,
    unresolved: bool = False,
) -> None:
    _add_paragraph(doc, title, HEADING1_STYLE)
    items: List[Any] = []
    for entry in entries:
        values = entry.get(field_name, []) or []
        if isinstance(values, list):
            items.extend(values)
        elif values:
            items.append(values)

    if unresolved:
        unresolved_items: List[Any] = []
        for entry in entries:
            values = entry.get("unresolved", []) or []
            if isinstance(values, list):
                unresolved_items.extend(values)
            elif values:
                unresolved_items.append(values)
        unresolved_items = _deduplicate(unresolved_items)
        unresolved_markers = {_canonical_key(item) for item in unresolved_items}
        progress_items = [
            item
            for item in _deduplicate(items)
            if _canonical_key(item) not in unresolved_markers
        ]
        if not progress_items and not unresolved_items:
            _add_item_paragraph(doc, "无")
            return
        for item in progress_items:
            _add_item_paragraph(doc, item)
        for item in unresolved_items:
            _add_item_paragraph(doc, item, unresolved=True)
        return

    items = _deduplicate(items)
    if not items:
        _add_item_paragraph(doc, "无")
        return
    for item in items:
        _add_item_paragraph(doc, item)


def generate_report(
    root: Path,
    week_of: date,
    template_path: Path,
    output_path: Optional[Path] = None,
) -> Path:
    root = Path(root)
    template_path = Path(template_path)
    entries = weekly_store.list_entries(root, week_of)
    if not entries:
        raise LookupError("no entries for week %s" % week_of.isoformat())

    target_path = Path(output_path) if output_path is not None else _default_output_path(root, week_of)
    doc = Document(str(template_path))

    _add_paragraph(doc, _week_range_label(week_of)[2], TITLE_STYLE)
    _render_daily_section(doc, entries)
    _render_flat_section(doc, "2. 阅读文献情况", entries, "literature")
    _render_flat_section(
        doc,
        "3. 研究进展情况",
        entries,
        "research_progress",
        unresolved=True,
    )
    _render_flat_section(doc, "4. 科研成果情况", entries, "research_outputs")

    return _atomic_save_document(doc, target_path)


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate weekly report docx files")
    commands = parser.add_subparsers(dest="command", required=True)

    create = commands.add_parser("create-template")
    create.add_argument("--output", required=True, type=Path)

    generate = commands.add_parser("generate")
    generate.add_argument("--root", required=True, type=Path)
    generate.add_argument("--week-of", required=True)
    generate.add_argument("--template", required=True, type=Path)
    generate.add_argument("--output", type=Path)
    return parser


def main(argv: Optional[Sequence[str]] = None) -> int:
    parser = _build_parser()
    arguments = parser.parse_args(argv)
    try:
        if arguments.command == "create-template":
            output = create_template(arguments.output)
            print(str(output.resolve()))
            return 0
        if arguments.command == "generate":
            week_of = weekly_store.normalize_date(arguments.week_of)
            output = generate_report(
                arguments.root,
                week_of,
                arguments.template,
                output_path=arguments.output,
            )
            print(str(output.resolve()))
            return 0
        raise ValueError("unknown command: %s" % arguments.command)
    except Exception as exc:
        print("error: %s" % exc, file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
